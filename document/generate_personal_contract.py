"""
生成个人项目技术服务合同 DOCX（使用 python-docx）
运行: python document/generate_personal_contract.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUT_DIR = os.path.join(os.path.dirname(__file__) or ".", "")
OUTPUT = os.path.join(OUT_DIR, "个人项目技术服务合同_水泥配比计算系统.docx")

# ── 颜色 ──
DARK_BLUE = RGBColor(0x1E, 0x3C, 0x72)
MID_BLUE = RGBColor(0x2A, 0x52, 0x98)
LIGHT_BLUE = RGBColor(0xE8, 0xF0, 0xFB)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x88, 0x88, 0x88)
TEXT_COLOR = RGBColor(0x33, 0x33, 0x33)

FONT_NAME = "微软雅黑"


# ── 工具函数 ──

def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_run(paragraph, text, bold=False, size=11, color=None, font_name=FONT_NAME):
    """向段落添加 run"""
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    return run


def add_paragraph(doc, text, bold=False, size=11, align=None, color=None,
                  space_before=3, space_after=3, font_name=FONT_NAME):
    """添加段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    add_run(p, text, bold=bold, size=size, color=color, font_name=font_name)
    return p


def add_bullet(doc, text, size=11, indent_left=Cm(0.72), color=None):
    """添加带圆点符号的段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = indent_left
    p.paragraph_format.first_line_indent = Cm(-0.36)
    add_run(p, "\u25cf ", size=size - 2, color=DARK_BLUE, font_name=FONT_NAME)
    c = color if color else TEXT_COLOR
    add_run(p, text, size=size, color=c, font_name=FONT_NAME)
    return p


def add_heading_styled(doc, text, level=1):
    """添加带样式的标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        if level == 1:
            run.font.color.rgb = DARK_BLUE
            run.font.size = Pt(16)
        elif level == 2:
            run.font.color.rgb = MID_BLUE
            run.font.size = Pt(13)
    return h


def add_table(doc, headers, rows, col_widths_cm=None):
    """添加带样式的表格，col_widths_cm 为各列宽度（cm）列表"""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        run.font.color.rgb = WHITE
        set_cell_shading(cell, "1E3C72")

    # 数据行
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            if isinstance(val, dict):
                align_val = val.get("align", WD_ALIGN_PARAGRAPH.LEFT)
                p.alignment = align_val
                text = val.get("text", "")
                run = p.add_run(text)
                if val.get("bold"):
                    run.bold = True
                if val.get("color"):
                    run.font.color.rgb = val["color"]
                if val.get("fill"):
                    set_cell_shading(cell, val["fill"])
                if val.get("size"):
                    run.font.size = Pt(val["size"])
            elif isinstance(val, list):
                # 多行文本
                for line_idx, line_text in enumerate(val):
                    if line_idx == 0:
                        run = p.add_run(line_text)
                    else:
                        np = cell.add_paragraph()
                        np.alignment = p.alignment
                        run = np.add_run(line_text)
                set_cell_shading(cell, "E8F0FB")
            else:
                run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = FONT_NAME
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # 列宽
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def empty_para(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)


def page_break(doc):
    doc.add_page_break()


def divider(doc):
    """添加装饰分割线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="E8F0FB"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


# ═══════════════════════════════════════
# 构建文档
# ═══════════════════════════════════════

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

# ── 页眉 ──
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hrun = hp.add_run("个人项目技术服务合同 · 水泥配比计算系统")
hrun.font.size = Pt(8)
hrun.font.name = FONT_NAME
hrun._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
hrun.font.color.rgb = GRAY
pPr = hp._element.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="2A5298"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

# ── 页脚 ──
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fPPr = fp._element.get_or_add_pPr()
fBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:top w:val="single" w:sz="2" w:space="1" w:color="CCCCCC"/>'
    f'</w:pBdr>'
)
fPPr.append(fBdr)
frun = fp.add_run("— ")
frun.font.size = Pt(8)
fP1 = fp.add_run()
# 页码域
fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
fP1._element.append(fldChar_begin)
fP1._element.append(instrText)
fP1._element.append(fldChar_end)
frun2 = fp.add_run(" —")
frun2.font.size = Pt(8)

# ═══════════════════════════════════════
# 封  面
# ═══════════════════════════════════════

for _ in range(8):
    empty_para(doc)

add_paragraph(doc, "个人项目技术服务合同", bold=True, size=22,
              align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE, space_after=12)
add_paragraph(doc, "水泥配比计算系统", bold=True, size=16,
              align=WD_ALIGN_PARAGRAPH.CENTER, color=MID_BLUE, space_after=24)
add_paragraph(doc, "混凝土配合比设计与试配调整平台", size=11,
              align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_after=36)
divider(doc)
add_paragraph(doc, "合同编号：SC-P-2026-0604", size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)
add_paragraph(doc, "签订日期：________ 年 ________ 月 ________ 日", size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)
add_paragraph(doc, "签订地点：____________________", size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)
page_break(doc)

# ═══════════════════════════════════════
# 第一条  合同双方
# ═══════════════════════════════════════
add_heading_styled(doc, "第一条  合同双方")
add_paragraph(doc, "本合同由以下双方在平等自愿、协商一致的基础上签订：", size=10)
empty_para(doc)

add_table(doc, ["项目", "信息"], [
    [{"text": "甲方（委托方）", "bold": True, "fill": "E8F0FB"},
     ["姓    名：________________________",
      "身份证号：________________________",
      "联系电话：________________________",
      "电子邮箱：________________________",
      "通讯地址：________________________"]],
    [{"text": "乙方（开发方）", "bold": True, "fill": "E8F0FB"},
     ["姓    名：________________________",
      "身份证号：________________________",
      "联系电话：________________________",
      "电子邮箱：________________________",
      "通讯地址：________________________"]],
], col_widths_cm=[3.0, 13.5])

add_paragraph(doc, "双方确认上述信息真实有效，并作为合同履行期间的送达地址。如一方信息变更，应在 3 日内书面通知另一方。",
              size=8, color=GRAY)
empty_para(doc)

# ═══════════════════════════════════════
# 第二条  项目概述
# ═══════════════════════════════════════
add_heading_styled(doc, "第二条  项目概述")
add_heading_styled(doc, "2.1  项目名称", level=2)
add_paragraph(doc, "水泥配比计算系统（混凝土配合比设计与试配调整平台）。")
add_heading_styled(doc, "2.2  项目定位", level=2)
add_paragraph(doc, "面向混凝土行业的配合比全流程数字化计算与试配调整工具，支持高性能混凝土（HPC）与超高性能混凝土（UHPC）的配合比设计与优化。")
add_heading_styled(doc, "2.3  技术方案", level=2)
add_paragraph(doc, "采用前后端分离架构，前端基于 Vue 3 + Element Plus，后端基于 Python FastAPI + SQLite 数据库。支持 Windows / Linux 跨平台部署，提供 Web 版与桌面版（Electron）两种交付形态。")
add_heading_styled(doc, "2.4  项目现状", level=2)
add_paragraph(doc, "本项目已完成核心功能开发，当前处于功能完善与部署交付阶段。本合同约定范围为：现有功能的交付部署、使用培训、免费维护期内的技术支持与 Bug 修复。")
empty_para(doc)

# ═══════════════════════════════════════
# 第三条  服务内容与交付清单
# ═══════════════════════════════════════
add_heading_styled(doc, "第三条  服务内容与交付清单")
add_paragraph(doc, "乙方按照以下功能清单完成系统的交付、部署与培训服务：")
empty_para(doc)

feature_headers = ["序号", "功能模块", "功能说明"]
features = [
    ["1", "系统基础", "用户认证与鉴权、角色权限控制、密码安全策略"],
    ["2", "首页仪表盘", "统计卡片、快捷操作面板（5格）、最近项目/记录展示"],
    ["3", "项目管理", "项目 CRUD、项目详情、关联配比记录管理、软删除回收站"],
    ["4", "HPC 配比计算", "水胶比（鲍罗米公式）、砂率选取、骨料用量（体积法）、胶凝材料（浆体体积法）、水与外加剂"],
    ["5", "UHPC 配比计算", "水胶比、砂胶比与钢纤维、胶凝材料比例（修正安德森模型）"],
    ["6", "试配调整", "HPC/UHPC 三阶段调整：工作性→强度回归→表观密度校正→实验室配合比"],
    ["7", "计算书", "项目选择→记录浏览→HTML 计算书→浏览器打印/导出 PDF"],
    ["8", "全部记录", "12 列数据表格、名称/项目搜索、HPC/UHPC 筛选、历史记录载入"],
    ["9", "用户管理", "管理员专属：用户 CRUD、密码重置"],
    ["10", "个人中心", "个人资料修改、密码修改（原密码验证）"],
    ["11", "回收站", "软删除管理、类型筛选/搜索、项目-记录关联恢复、物理删除确认"],
    ["12", "系统部署", "Web 版（Nginx 部署）、桌面版（Electron + 后端守护进程）、数据库自动初始化"],
]

rows = []
for i, f in enumerate(features):
    fill = "E8F0FB" if i % 2 == 0 else None
    rows.append([
        {"text": f[0], "align": WD_ALIGN_PARAGRAPH.CENTER, "fill": fill},
        {"text": f[1], "bold": True, "fill": fill},
        {"text": f[2], "fill": fill},
    ])

add_table(doc, feature_headers, rows, col_widths_cm=[1.0, 3.0, 12.5])
add_paragraph(doc, "以上共计 12 项功能模块。乙方确保各项功能正常运行并满足甲方使用需求。", size=9, color=GRAY)
empty_para(doc)

# ═══════════════════════════════════════
# 第四条  交付物清单
# ═══════════════════════════════════════
add_heading_styled(doc, "第四条  交付物清单")
empty_para(doc)

deliverable_headers = ["序号", "交付物名称", "交付形式", "交付时间"]
deliverables = [
    ["1", "Web 版部署包（前端 dist + 后端可执行文件 + Nginx 配置）", "电子文件", "____年____月____日"],
    ["2", "桌面版安装包（Windows / Linux）", "安装程序", "____年____月____日"],
    ["3", "《软件操作说明书》", "电子版（PDF）", "____年____月____日"],
    ["4", "《功能验证清单》", "电子版（PDF）", "____年____月____日"],
    ["5", "源代码（如有约定交付）", "电子文件", "____年____月____日"],
    ["6", "本合同", "纸质/电子各一份", "签订当日"],
]

drows = []
for i, d in enumerate(deliverables):
    fill = "E8F0FB" if i % 2 == 0 else None
    drows.append([
        {"text": d[0], "align": WD_ALIGN_PARAGRAPH.CENTER, "fill": fill},
        {"text": d[1], "fill": fill},
        {"text": d[2], "align": WD_ALIGN_PARAGRAPH.CENTER, "fill": fill},
        {"text": d[3], "align": WD_ALIGN_PARAGRAPH.CENTER, "fill": fill},
    ])

add_table(doc, deliverable_headers, drows, col_widths_cm=[1.0, 7.5, 3.0, 5.0])
empty_para(doc)

# ═══════════════════════════════════════
# 第五条  费用与付款方式
# ═══════════════════════════════════════
add_heading_styled(doc, "第五条  费用与付款方式")
add_heading_styled(doc, "5.1  项目总费用", level=2)
add_paragraph(doc, "本项目总费用为人民币 ________________ 元整（¥_______________）。")
add_paragraph(doc, "费用包含：系统交付部署、使用培训、操作手册编写、免费维护期内的技术支持与 Bug 修复。")
add_paragraph(doc, "费用不包含：服务器/域名/SSL 证书等第三方资源费用、超出本合同约定范围的定制开发。")
add_heading_styled(doc, "5.2  付款计划", level=2)
empty_para(doc)

add_table(doc, ["付款节点", "付款比例", "金额（元）"], [
    ["合同签订后 3 个工作日内（预付款）", {"text": "____%", "align": WD_ALIGN_PARAGRAPH.CENTER},
     {"text": "_______________", "align": WD_ALIGN_PARAGRAPH.CENTER}],
    [{"text": "系统部署完成并通过验收后 5 个工作日内", "fill": "E8F0FB"},
     {"text": "____%", "align": WD_ALIGN_PARAGRAPH.CENTER, "fill": "E8F0FB"},
     {"text": "_______________", "align": WD_ALIGN_PARAGRAPH.CENTER, "fill": "E8F0FB"}],
], col_widths_cm=[8.5, 3.0, 5.0])

add_paragraph(doc, "甲方应通过银行转账方式支付至乙方指定账户。乙方收款账户信息以书面（含电子邮件）形式单独提供。",
              size=9, color=GRAY)
empty_para(doc)

# ═══════════════════════════════════════
# 第六条  交付周期
# ═══════════════════════════════════════
add_heading_styled(doc, "第六条  交付周期")
add_paragraph(doc, "6.1  本合同签订后，乙方应在 ________ 个工作日内完成系统部署、功能验证、使用培训及全部交付物的移交。")
add_paragraph(doc, "6.2  如因甲方原因（如未能及时提供服务器环境、未能及时安排验收等）导致交付延迟的，交付时间相应顺延。")
add_paragraph(doc, "6.3  如遇不可抗力（自然灾害、政策变化、重大疫情等），双方协商调整交付时间。")
empty_para(doc)

# ═══════════════════════════════════════
# 第七条  验收标准与方式
# ═══════════════════════════════════════
add_heading_styled(doc, "第七条  验收标准与方式")
add_paragraph(doc, "7.1  验收依据：本合同第三条所列功能清单及双方确认的《功能验证清单》。")
add_paragraph(doc, "7.2  验收流程：")
add_bullet(doc, "甲方在乙方部署完成后 5 个工作日内，依据功能清单逐项进行验收测试；")
add_bullet(doc, "甲方填写验收结论并签字（或电子邮件）确认；")
add_bullet(doc, "如发现功能缺陷，乙方应在 ________ 个工作日内完成修复并提请复验。")
add_paragraph(doc, "7.3  验收合格后，双方确认《项目验收确认书》，系统进入免费维护期。")
add_paragraph(doc, "7.4  甲方无正当理由在部署完成后超过 15 个工作日未组织验收或未提出书面异议的，视为验收通过。",
              size=9, color=GRAY)
empty_para(doc)

# ═══════════════════════════════════════
# 第八条  维护与技术支持
# ═══════════════════════════════════════
add_heading_styled(doc, "第八条  维护与技术支持")
add_paragraph(doc, "8.1  免费维护期：自项目验收合格之日起 ________ 个月。")
add_paragraph(doc, "8.2  免费维护期内，乙方提供以下服务：")
add_bullet(doc, "修复影响正常使用的程序 Bug（乙方在收到通知后 ________ 个工作日内响应）；")
add_bullet(doc, "提供远程技术支持（电话/微信/远程桌面/邮件等）；")
add_bullet(doc, "因操作系统或浏览器大版本升级导致的兼容性问题修复。")
add_paragraph(doc, "8.3  免费维护期后，双方可协商签订年度维护合同，年维护费建议不超过项目总费用的 15%。")
add_paragraph(doc, "8.4  以下情况不在免费维护范围内：")
add_bullet(doc, "甲方或第三方自行修改源代码、数据库导致的故障；")
add_bullet(doc, "超出原定功能范围的新需求开发；")
add_bullet(doc, "因硬件故障、网络攻击、病毒等不可抗力导致的问题。")
empty_para(doc)

# ═══════════════════════════════════════
# 第九条  知识产权
# ═══════════════════════════════════════
add_heading_styled(doc, "第九条  知识产权")
add_paragraph(doc, "9.1  本项目产生的全部源代码及文档的知识产权归属，双方约定如下（勾选一项）：")
add_paragraph(doc, "    □ 甲方享有全部知识产权，乙方向甲方完整交付全部源代码；")
add_paragraph(doc, "    □ 乙方保留知识产权，甲方享有本项目的永久使用权和复制权；")
add_paragraph(doc, "    □ 双方共有知识产权，具体权利划分另行书面约定。")
add_paragraph(doc, "9.2  无论知识产权如何归属，乙方均有权在个人作品集中展示本项目（不得泄露甲方商业秘密）。")
add_paragraph(doc, "9.3  乙方保证所使用的第三方库和工具均为开源或已获得合法授权，不侵犯第三方知识产权。")
empty_para(doc)

# ═══════════════════════════════════════
# 第十条  保密条款
# ═══════════════════════════════════════
add_heading_styled(doc, "第十条  保密条款")
add_paragraph(doc, "10.1  双方应对在合同履行过程中知悉的对方商业秘密、技术秘密及非公开信息承担保密义务。")
add_paragraph(doc, "10.2  保密义务不因合同履行完毕或终止而失效，保密期限为合同终止后 ________ 年。")
add_paragraph(doc, "10.3  任何一方违反保密义务给对方造成损失的，应承担赔偿责任。")
empty_para(doc)

# ═══════════════════════════════════════
# 第十一条  违约责任
# ═══════════════════════════════════════
add_heading_styled(doc, "第十一条  违约责任")
add_paragraph(doc, "11.1  任何一方违反本合同约定，应承担相应的违约责任，并赔偿对方因此遭受的直接经济损失。")
add_paragraph(doc, "11.2  因不可抗力导致合同无法履行的，双方互不承担违约责任，但应及时通知对方并协商后续处理方案。")
add_paragraph(doc, "11.3  如甲方逾期支付款项，每逾期一日应按未付款项的 0.05% 向乙方支付违约金。")
add_paragraph(doc, "11.4  如乙方逾期交付，每逾期一日应按项目总费用的 0.05% 向甲方支付违约金，但违约金总额不超过项目总费用的 10%。")
empty_para(doc)

# ═══════════════════════════════════════
# 第十二条  争议解决
# ═══════════════════════════════════════
add_heading_styled(doc, "第十二条  争议解决")
add_paragraph(doc, "12.1  本合同在履行过程中如发生争议，双方应首先友好协商解决。")
add_paragraph(doc, "12.2  协商不成的，任何一方可向乙方住所地有管辖权的人民法院提起诉讼。")
empty_para(doc)

# ═══════════════════════════════════════
# 第十三条  其他约定
# ═══════════════════════════════════════
add_heading_styled(doc, "第十三条  其他约定")
add_paragraph(doc, "13.1  本合同一式两份，甲乙双方各执一份，具有同等法律效力。")
add_paragraph(doc, "13.2  本合同未尽事宜，由双方协商后签订补充协议，补充协议与本合同具有同等法律效力。")
add_paragraph(doc, "13.3  本合同自双方签字之日起生效。")
empty_para(doc, 4)

# ═══════════════════════════════════════
# 签章区
# ═══════════════════════════════════════
divider(doc)
empty_para(doc)

add_table(doc, ["甲方（委托方）", "乙方（开发方）"], [
    ["签字：\n\n\n\n\n日期：________ 年 ________ 月 ________ 日",
     "签字：\n\n\n\n\n日期：________ 年 ________ 月 ________ 日"],
], col_widths_cm=[8.25, 8.25])

# ── 保存 ──
doc.save(OUTPUT)
print(f"✅ 合同已生成: {OUTPUT}")
print(f"   文件大小: {os.path.getsize(OUTPUT) / 1024:.1f} KB")
