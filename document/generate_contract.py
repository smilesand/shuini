"""
生成项目合同 DOCX（使用 python-docx）
运行: python document/generate_contract.py
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUT_DIR = os.path.join(os.path.dirname(__file__) or ".", "")
OUTPUT = os.path.join(OUT_DIR, "项目合同_水泥配比计算系统.docx")

# ── 颜色 ──
DARK_BLUE = RGBColor(0x1E, 0x3C, 0x72)
MID_BLUE = RGBColor(0x2A, 0x52, 0x98)
LIGHT_BLUE = RGBColor(0xE8, 0xF0, 0xFB)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x88, 0x88, 0x88)
BLACK = RGBColor(0x33, 0x33, 0x33)

# ── 工具函数 ──

def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_paragraph(doc, text, bold=False, size=11, align=None, color=None, space_before=3, space_after=3, font_name="微软雅黑"):
    """添加段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    return p

def add_heading_styled(doc, text, level=1):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        if level == 1:
            run.font.color.rgb = DARK_BLUE
            run.font.size = Pt(16)
        elif level == 2:
            run.font.color.rgb = MID_BLUE
            run.font.size = Pt(13)
    return h

def add_table(doc, headers, rows, col_widths=None):
    """添加带样式的表格"""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        run.font.color.rgb = WHITE
        set_cell_shading(cell, "1E3C72")

    # 数据行
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            if isinstance(val, dict):
                p.alignment = val.get("align", WD_ALIGN_PARAGRAPH.LEFT)
                run = p.add_run(val.get("text", ""))
                if val.get("bold"):
                    run.bold = True
                if val.get("color"):
                    run.font.color.rgb = val["color"]
                if val.get("fill"):
                    set_cell_shading(cell, val["fill"])
            else:
                run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def empty_para(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

def page_break(doc):
    doc.add_page_break()


# ═══════════════════════════════════════
# 构建文档
# ═══════════════════════════════════════

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

# ── 页眉 ──
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hrun = hp.add_run("项目合同 · 水泥配比计算系统")
hrun.font.size = Pt(8)
hrun.font.name = "微软雅黑"
hrun._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
hrun.font.color.rgb = GRAY
# 页眉底部边框
pPr = hp._element.get_or_add_pPr()
pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="2A5298"/></w:pBdr>')
pPr.append(pBdr)

# ── 页脚 ──
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
# 页脚顶部边框
fPPr = fp._element.get_or_add_pPr()
fBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="2" w:space="1" w:color="CCCCCC"/></w:pBdr>')
fPPr.append(fBdr)
frun = fp.add_run("— ")
frun.font.size = Pt(8)
fP1 = fp.add_run()
# 页码域
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
fP1._element.append(fldChar1)
fP1._element.append(instrText)
fP1._element.append(fldChar2)
frun2 = fp.add_run(" —")
frun2.font.size = Pt(8)

# ═══════════════════════════════════════
# 封 面
# ═══════════════════════════════════════

for _ in range(8):
    empty_para(doc, 1)

add_paragraph(doc, "项  目  合  同", bold=True, size=24, align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE, space_after=12)
add_paragraph(doc, "水泥配比计算系统", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, color=MID_BLUE, space_after=24)
add_paragraph(doc, "混凝土配合比设计与试配调整平台", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_after=36)
add_paragraph(doc, "合同编号：SC-2026-0602", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)
add_paragraph(doc, "签订日期：________ 年 ________ 月 ________ 日", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)
add_paragraph(doc, "签订地点：____________________", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)
page_break(doc)

# ═══════════════════════════════════════
# 第一条 合同双方
# ═══════════════════════════════════════
add_heading_styled(doc, "第一条  合同双方")
add_table(doc, ["项目", "信息"], [
    [{"text": "甲方（委托方）", "bold": True, "fill": "E8F0FB"},
     "公司名称：________________________\n联系人：________________________\n联系电话：________________________\n电子邮箱：________________________"],
    [{"text": "乙方（开发方）", "bold": True, "fill": "E8F0FB"},
     "姓名/公司：________________________\n联系人：________________________\n联系电话：________________________\n电子邮箱：________________________"],
], col_widths=[3.5, 13.0])
empty_para(doc)

# ═══════════════════════════════════════
# 第二条 项目概述
# ═══════════════════════════════════════
add_heading_styled(doc, "第二条  项目概述")
add_paragraph(doc, "2.1  项目名称：水泥配比计算系统（混凝土配合比设计与试配调整平台）。")
add_paragraph(doc, "2.2  项目定位：面向中国中车风电混塔项目的高性能混凝土（HPC）与超高性能混凝土（UHPC）配合比全流程数字化设计与试配调整系统。")
add_paragraph(doc, "2.3  技术路线：前后端分离 Web 应用，前端基于 Vue 3 + Element Plus，后端基于 Python FastAPI + SQLite，支持 Windows / Linux 跨平台部署。")
empty_para(doc)

# ═══════════════════════════════════════
# 第三条 功能范围
# ═══════════════════════════════════════
add_heading_styled(doc, "第三条  功能范围与交付清单")
add_paragraph(doc, "乙方按照以下功能清单完成全部软件功能的设计、开发、测试与交付：")
empty_para(doc)

feature_headers = ["序号", "功能模块", "功能名称", "功能说明"]
features = [
    ["1", "系统基础", "用户认证与鉴权", "JWT Token 登录认证，过期自动跳转；前端路由守卫 + 后端二次鉴权"],
    ["2", "系统基础", "角色与权限控制", "管理员/普通用户双角色，菜单级权限隔离，API 层二次校验"],
    ["3", "系统基础", "密码安全策略", "SHA-256 加盐哈希存储；首次登录强制提示改密；管理员可重置用户密码"],
    ["4", "首页总览", "首页仪表盘", "统计卡片 + 快捷操作面板(5格) + 最近项目5条 + 最近记录10条"],
    ["5", "项目管理", "项目 CRUD", "项目列表分页搜索；新建/编辑/删除项目；软删除进入回收站"],
    ["6", "项目管理", "项目详情", "项目信息展示 + 关联配比记录表格(12列) + 新建/载入/删除记录"],
    ["7", "HPC配比计算", "水胶比页签", "鲍罗米公式自动计算 fcu,0 和 W/B；fb 自动估算；文件上传回归拟合 αa/αb"],
    ["8", "HPC配比计算", "砂率选取页签", "标准砂率参考表(强度等级×粒径)，行列交叉高亮，手动输入确认"],
    ["9", "HPC配比计算", "骨料用量页签", "体积法计算 mg/ms；Vg 参考表；粗细骨料密度输入"],
    ["10", "HPC配比计算", "胶凝材料页签", "浆体体积法；4种掺合料(粉煤灰/矿粉/微珠/硅灰)质量分数与密度输入"],
    ["11", "HPC配比计算", "水与外加剂页签", "mw 计算；外加剂百分比计算 ma；最终配比汇总表"],
    ["12", "UHPC配比计算", "水胶比页签", "推荐强度等级-W/B 对照表，点击自动带入；外加剂默认值自动补入"],
    ["13", "UHPC配比计算", "砂胶比与钢纤维", "砂胶比体积参数计算；钢纤维体积掺量与质量计算"],
    ["14", "UHPC配比计算", "胶凝材料比例", "修正安德森模型(D10/D50/D90)；体积比例→质量比例→每方用量"],
    ["15", "HPC试配调整", "工作性-强度-校正", "三阶段：Δ调整→强度回归分析→表观密度校正→实验室配合比"],
    ["16", "UHPC试配调整", "工作性-强度-校正", "三阶段：S/B/α调整→+Δ/-Δ回归→三种基准校正→实验室配合比"],
    ["17", "计算书", "计算书生成", "项目选择→记录浏览→HTML计算书→浏览器打印/导出PDF"],
    ["18", "全部记录", "记录管理", "12列数据表格；名称/项目搜索；HPC/UHPC筛选；载入历史记录"],
    ["19", "回收站", "数据恢复与清空", "软删除管理；类型筛选/搜索；项目-记录关联恢复；物理删除确认"],
    ["20", "用户管理", "用户 CRUD", "管理员专属；用户列表搜索；新增(默认密码)；密码重置；删除"],
    ["21", "个人中心", "资料与密码", "邮箱/手机号修改；原密码验证→新密码确认→自动退出重新登录"],
    ["22", "系统部署", "前后端打包部署", "Vite 生产构建 + Nginx 部署；PyInstaller 后端打包；数据库自动初始化"],
    ["23", "系统部署", "数据库与日志", "SQLite 自动建表；默认管理员初始化；文件+控制台日志按日轮转"],
]

rows = []
for i, f in enumerate(features):
    fill = "E8F0FB" if i % 2 == 0 else None
    rows.append([
        {"text": f[0], "align": WD_ALIGN_PARAGRAPH.CENTER, "fill": fill},
        {"text": f[1], "fill": fill},
        {"text": f[2], "bold": True, "fill": fill},
        {"text": f[3], "fill": fill},
    ])

add_table(doc, feature_headers, rows, col_widths=[1.0, 2.5, 3.0, 10.0])
add_paragraph(doc, "以上共计 23 项功能模块，覆盖全部业务场景。", size=9, color=GRAY, space_before=6)
empty_para(doc)

# ═══════════════════════════════════════
# 第四条 项目报价
# ═══════════════════════════════════════
add_heading_styled(doc, "第四条  项目报价")
add_paragraph(doc, "4.1  项目总报价：人民币 ________________ 元整（¥________________）。")
add_paragraph(doc, "4.2  报价包含：全部功能模块的设计与开发、前端构建与 Nginx 部署、后端打包与运行环境配置、数据库初始化与管理账号预置、部署上线与基础使用培训、验收交付后的免费 Bug 修复与维护支持。")
add_paragraph(doc, "4.3  报价不包含：服务器/域名/SSL证书等第三方资源费用、超出约定范围的定制开发。")
empty_para(doc)

add_table(doc, ["费用项目", "内容说明", "金额（元）", "备注"], [
    ["软件开发费", "全部功能模块的设计、开发与测试", "____________", "核心交付"],
    [{"text": "部署实施费", "fill": "E8F0FB"}, {"text": "前端 Nginx 部署 + 后端打包 + 环境配置", "fill": "E8F0FB"},
     {"text": "____________", "align": WD_ALIGN_PARAGRAPH.CENTER, "fill": "E8F0FB"}, {"text": "一次性", "fill": "E8F0FB"}],
    ["培训与文档费", "操作说明书 + 验收清单 + 使用培训", "____________", "交付物"],
    [{"text": "首年维护费", "fill": "E8F0FB"}, {"text": "Bug修复 + 小版本更新 + 技术支持", "fill": "E8F0FB"},
     {"text": "____________", "align": WD_ALIGN_PARAGRAPH.CENTER, "fill": "E8F0FB"}, {"text": "每年", "fill": "E8F0FB"}],
    [{"text": "合计", "bold": True}, "", {"text": "____________", "bold": True, "align": WD_ALIGN_PARAGRAPH.CENTER, "color": DARK_BLUE}, {"text": "含税", "bold": True}],
], col_widths=[3.0, 5.0, 4.0, 4.5])
empty_para(doc)

# ═══════════════════════════════════════
# 第五条 开发周期
# ═══════════════════════════════════════
add_heading_styled(doc, "第五条  开发周期与交付计划")
add_paragraph(doc, "5.1  项目总工期：自合同签订之日起 ________ 个工作日。")
add_paragraph(doc, "5.2  里程碑节点：")
empty_para(doc)

add_table(doc, ["阶段", "交付物", "完成时间"], [
    ["第一阶段：基础框架搭建", "登录认证、项目管理、数据库与API基础", {"text": "____年____月____日", "align": WD_ALIGN_PARAGRAPH.CENTER}],
    [{"text": "第二阶段：核心计算模块", "fill": "E8F0FB"}, {"text": "HPC/UHPC配比计算与试配调整", "fill": "E8F0FB"},
     {"text": "____年____月____日", "align": WD_ALIGN_PARAGRAPH.CENTER, "fill": "E8F0FB"}],
    ["第三阶段：辅助功能与联调", "计算书、回收站、用户管理、全部记录", {"text": "____年____月____日", "align": WD_ALIGN_PARAGRAPH.CENTER}],
    [{"text": "第四阶段：部署上线与验收", "fill": "E8F0FB"}, {"text": "生产构建、部署配置、验收测试、文档交付", "fill": "E8F0FB"},
     {"text": "____年____月____日", "align": WD_ALIGN_PARAGRAPH.CENTER, "fill": "E8F0FB"}],
], col_widths=[5.5, 7.0, 4.0])
empty_para(doc)

# ═══════════════════════════════════════
# 第六条 部署与交付
# ═══════════════════════════════════════
add_heading_styled(doc, "第六条  部署与交付")
add_paragraph(doc, "6.1  乙方负责完成以下部署工作：")
add_paragraph(doc, "（1）前端资源构建（npm run build:prod）并部署至甲方指定的 Nginx 服务器；")
add_paragraph(doc, "（2）后端使用 PyInstaller 打包为独立可执行文件，配置为系统服务或开机自启；")
add_paragraph(doc, "（3）数据库首次启动自动初始化，创建全部表结构及默认管理员账号；")
add_paragraph(doc, "（4）配置 Nginx 反向代理规则，实现前后端统一入口与 API 代理。")
add_paragraph(doc, "6.2  交付物清单：")
add_paragraph(doc, "（1）前端构建产物（dist 目录）；")
add_paragraph(doc, "（2）后端可执行文件（main）及配置文件；")
add_paragraph(doc, "（3）《软件操作说明书》一份（电子版）；")
add_paragraph(doc, "（4）《功能清单及验收清单》一份（电子版）；")
add_paragraph(doc, "（5）本项目合同一份。")
empty_para(doc)

# ═══════════════════════════════════════
# 第七条 维护与技术支持
# ═══════════════════════════════════════
add_heading_styled(doc, "第七条  维护与技术支持")
add_paragraph(doc, "7.1  免费维护期：自项目验收合格之日起 ________ 个月。")
add_paragraph(doc, "7.2  免费维护期内，乙方提供以下服务：")
add_paragraph(doc, "（1）修复影响正常使用的程序 Bug（乙方在收到通知后 ________ 个工作日内响应）；")
add_paragraph(doc, "（2）提供远程技术支持（电话/即时通讯/远程桌面）；")
add_paragraph(doc, "（3）因操作系统或浏览器大版本升级导致的兼容性问题修复。")
add_paragraph(doc, "7.3  免费维护期后，双方可另行签订年度维护合同。")
add_paragraph(doc, "7.4  以下情况不在免费维护范围内：")
add_paragraph(doc, "（1）甲方自行修改代码或数据库导致的故障；")
add_paragraph(doc, "（2）超出原定功能范围的新需求开发；")
add_paragraph(doc, "（3）因硬件故障、网络攻击等不可抗力导致的问题。")
empty_para(doc)

# ═══════════════════════════════════════
# 第八条 付款方式
# ═══════════════════════════════════════
add_heading_styled(doc, "第八条  付款方式")
empty_para(doc)

add_table(doc, ["付款节点", "付款比例", "金额（元）"], [
    ["合同签订后 5 个工作日内（预付款）", {"text": "____%", "align": WD_ALIGN_PARAGRAPH.CENTER}, {"text": "____________", "align": WD_ALIGN_PARAGRAPH.CENTER}],
    [{"text": "核心功能开发完成并演示通过", "fill": "E8F0FB"}, {"text": "____%", "align": WD_ALIGN_PARAGRAPH.CENTER, "fill": "E8F0FB"},
     {"text": "____________", "align": WD_ALIGN_PARAGRAPH.CENTER, "fill": "E8F0FB"}],
    ["全部功能验收合格后 5 个工作日内", {"text": "____%", "align": WD_ALIGN_PARAGRAPH.CENTER}, {"text": "____________", "align": WD_ALIGN_PARAGRAPH.CENTER}],
], col_widths=[8.5, 3.0, 5.0])
add_paragraph(doc, "甲方通过银行转账方式支付至乙方指定账户。", size=9, color=GRAY, space_before=6)
empty_para(doc)

# ═══════════════════════════════════════
# 第九条 验收
# ═══════════════════════════════════════
add_heading_styled(doc, "第九条  验收标准与方式")
add_paragraph(doc, "9.1  验收依据：本合同第三条所列功能清单及双方确认的《功能清单及验收清单》。")
add_paragraph(doc, "9.2  验收方式：甲方逐项测试全部 46 项验收指标，填写验收结论并签字确认。")
add_paragraph(doc, "9.3  验收不合格的处理：如部分功能不符合约定，乙方应在 ________ 个工作日内完成整改并提请复验。")
add_paragraph(doc, "9.4  验收合格后，双方签署《项目验收确认书》，进入免费维护期。")
empty_para(doc)

# ═══════════════════════════════════════
# 第十条 知识产权
# ═══════════════════════════════════════
add_heading_styled(doc, "第十条  知识产权")
add_paragraph(doc, "10.1  本项目产生的全部源代码、文档、设计成果的知识产权归双方协商确定（□ 甲方所有 / □ 乙方所有 / □ 双方共有）。")
add_paragraph(doc, "10.2  乙方保证所使用的第三方库和工具均为开源或已获得合法授权，不侵犯第三方知识产权。")
empty_para(doc)

# ═══════════════════════════════════════
# 第十一条 违约责任
# ═══════════════════════════════════════
add_heading_styled(doc, "第十一条  违约责任")
add_paragraph(doc, "11.1  任何一方违反本合同约定，应承担相应的违约责任，并赔偿对方因此遭受的直接经济损失。")
add_paragraph(doc, "11.2  因不可抗力导致合同无法履行的，双方互不承担违约责任。")
empty_para(doc)

# ═══════════════════════════════════════
# 第十二条 其他
# ═══════════════════════════════════════
add_heading_styled(doc, "第十二条  其他约定")
add_paragraph(doc, "12.1  本合同一式两份，甲乙双方各执一份，具有同等法律效力。")
add_paragraph(doc, "12.2  本合同未尽事宜，由双方协商后签订补充协议。")
add_paragraph(doc, "12.3  本合同自双方签字（或盖章）之日起生效。")
empty_para(doc, 4)

# ═══════════════════════════════════════
# 签章区
# ═══════════════════════════════════════
add_table(doc, ["甲方（委托方）", "乙方（开发方）"], [
    ["签字/盖章：\n\n\n\n\n日期：________ 年 ________ 月 ________ 日",
     "签字/盖章：\n\n\n\n\n日期：________ 年 ________ 月 ________ 日"],
], col_widths=[8.25, 8.25])

# ── 保存 ──
doc.save(OUTPUT)
print(f"✅ 合同已生成: {OUTPUT}")
print(f"   文件大小: {os.path.getsize(OUTPUT) / 1024:.1f} KB")
